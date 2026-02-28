"""
export_service.py
Assembles a draft into a properly formatted manuscript DOCX.
Follows standard manuscript format conventions.
"""

import re
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from repositories.chapter_repo import get_chapters_for_draft
from repositories.draft_repo import get_draft_by_id
from repositories.manuscript_repo import get_manuscript_by_id
from services.permission_service import can_write


# ── Markdown stripping ────────────────────────────────────────────

def _parse_markdown_runs(text):
    """
    Parse a paragraph of markdown into a list of (text, bold, italic) tuples.
    Handles **bold**, *italic*, ***bold-italic***, and footnote refs.
    """
    # Footnote refs [^1] → (1)
    text = re.sub(r'\[\^([^\]]+)\]', r'(\1)', text)
    # Strip links [text](url) → text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

    runs = []
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|_(.+?)_)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            runs.append((text[last:m.start()], False, False))
        raw = m.group(0)
        if raw.startswith('***'):
            runs.append((m.group(2), True, True))
        elif raw.startswith('**'):
            runs.append((m.group(3), True, False))
        else:
            runs.append((m.group(4) or m.group(5), False, True))
        last = m.end()
    if last < len(text):
        runs.append((text[last:], False, False))
    return runs


def _extract_footnotes(md):
    """Pull footnote definitions out, return (body, {key: text})."""
    definitions = {}
    pattern = re.compile(r'^\[\^([^\]]+)\]:\s*(.+?)$', re.MULTILINE)
    for m in pattern.finditer(md):
        definitions[m.group(1)] = m.group(2).strip()
    body = pattern.sub('', md).strip()
    return body, definitions


def _split_paragraphs(md):
    """Split markdown body into paragraph chunks."""
    return [p.strip() for p in re.split(r'\n\n+', md) if p.strip()]


def _is_scene_break(text):
    return bool(re.match(r'^[\*\-#~\s]+$', text)) and len(text.strip()) <= 5


def _is_heading(text):
    return text.startswith('#')


def _strip_heading(text):
    return re.sub(r'^#{1,6}\s+', '', text)


# ── Document helpers ──────────────────────────────────────────────

def _set_para_spacing(para, space_before=0, space_after=0, line_spacing=None):
    fmt = para.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after  = Pt(space_after)
    if line_spacing:
        from docx.shared import Length
        fmt.line_spacing = line_spacing


def _add_header(doc, author_email, title, draft_name):
    """Add running header: Author / Title — Draft · Page #"""
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False

    # Clear default content
    for p in header.paragraphs:
        p.clear()

    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    short_title = title[:30] + '…' if len(title) > 30 else title
    run = p.add_run(f"{author_email} / {short_title} — {draft_name}  ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Page number field
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    run2 = p.add_run()
    run2._r.append(fld)

    instr = OxmlElement('w:instrText')
    instr.text = ' PAGE '
    run3 = p.add_run()
    run3._r.append(instr)

    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'end')
    run4 = p.add_run()
    run4._r.append(fld2)


def _add_title_block(doc, display_name, draft_name, author_email):
    """Standard manuscript title page block — top third of page."""
    # Push title down ~1/3 page with spacing
    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(display_name.upper())
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.font.bold = False

    by_p = doc.add_paragraph()
    by_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = by_p.add_run('by')
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = author_p.add_run(author_email)
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(12)

    draft_p = doc.add_paragraph()
    draft_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = draft_p.add_run(draft_name)
    r4.font.name = 'Times New Roman'
    r4.font.size = Pt(11)
    r4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def _add_chapter_heading(doc, title, chapter_num):
    """Centered chapter heading with page break."""
    doc.add_page_break()

    # Space down
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    heading_p = doc.add_paragraph()
    heading_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = heading_p.add_run(title)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.bold = False

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(24)


def _add_body_paragraph(doc, text, first_in_chapter=False):
    """Add a double-spaced body paragraph with proper indent."""
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after  = Pt(0)
    # Standard manuscript: double-spaced, half-inch first-line indent
    # except first paragraph after heading (no indent)
    from docx.shared import Length
    from docx.shared import Pt as PT
    fmt.line_spacing = Pt(24)  # double space at 12pt
    if not first_in_chapter:
        fmt.first_line_indent = Inches(0.5)

    runs = _parse_markdown_runs(text)
    for run_text, bold, italic in runs:
        if not run_text:
            continue
        r = p.add_run(run_text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.bold   = bold
        r.italic = italic

    return p


def _add_scene_break(doc):
    p = doc.add_paragraph('#')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.space_before = Pt(12)
    fmt.space_after  = Pt(12)
    fmt.line_spacing = Pt(24)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)


# ── Main builder ──────────────────────────────────────────────────

def build_docx(user_email, draft_id):
    draft = get_draft_by_id(draft_id)
    if not draft:
        raise ValueError("Draft not found")

    manuscript = get_manuscript_by_id(draft["manuscript_id"])
    if not manuscript:
        raise ValueError("Manuscript not found")

    is_author = can_write(user_email, manuscript_id=draft["manuscript_id"])
    chapters  = get_chapters_for_draft(draft_id, include_content=True)

    if not is_author:
        chapters = [c for c in chapters if c.get("status", "published") == "published"]
    if not chapters:
        raise ValueError("No chapters available to export")

    display_name = manuscript.get("display_name", "Untitled")
    draft_name   = draft.get("name", "Draft")

    # ── Build document ──
    doc = Document()

    # Page setup — US Letter, standard margins
    section = doc.sections[0]
    section.page_width   = Inches(8.5)
    section.page_height  = Inches(11)
    section.left_margin  = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin   = Inches(1)
    section.bottom_margin = Inches(1)

    # Running header
    _add_header(doc, user_email, display_name, draft_name)

    # Title block
    _add_title_block(doc, display_name, draft_name, user_email)

    # Chapters
    for chapter in chapters:
        title   = chapter.get("title", "Untitled")
        content = chapter.get("content", "")

        body, footnotes = _extract_footnotes(content)
        paragraphs = _split_paragraphs(body)

        _add_chapter_heading(doc, title, chapter.get("order", 0) + 1)

        first = True
        for para_text in paragraphs:
            if _is_scene_break(para_text):
                _add_scene_break(doc)
                first = False
                continue
            if _is_heading(para_text):
                # Subheadings within chapter — treat as scene break label
                p = doc.add_paragraph(_strip_heading(para_text))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = Pt(24)
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12)
                first = False
                continue

            _add_body_paragraph(doc, para_text, first_in_chapter=first)
            first = False

        # Footnotes at end of chapter if any
        if footnotes:
            sep = doc.add_paragraph()
            sep.paragraph_format.line_spacing = Pt(24)
            sep.add_run('─' * 20).font.name = 'Times New Roman'

            for key, fn_text in footnotes.items():
                fn_p = doc.add_paragraph()
                fn_p.paragraph_format.line_spacing = Pt(24)
                r = fn_p.add_run(f"({key}) {fn_text}")
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)

    # Save
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = re.sub(r'[^\w\s-]', '', display_name).strip().replace(' ', '_')
    safe_draft = re.sub(r'[^\w\s-]', '', draft_name).strip().replace(' ', '_')
    filename = f"{safe_title}-{safe_draft}.docx"

    return buf, filename
